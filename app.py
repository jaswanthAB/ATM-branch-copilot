"""
Branch Copilot — Streamlit Web App
Run with: streamlit run app.py
"""

import os, io
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from datetime import datetime
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Branch Copilot",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

*, html, body, [class*="css"] {
    font-family: 'Inter', sans-serif !important;
}

/* App background */
.stApp { background-color: #0d0f14; color: #c9cde0; }

/* Sidebar */
section[data-testid="stSidebar"] {
    background-color: #111318;
    border-right: 1px solid #1e2130;
}
section[data-testid="stSidebar"] .stRadio label {
    color: #8b90a8 !important;
    font-size: 0.85rem !important;
    padding: 6px 0;
}
section[data-testid="stSidebar"] .stRadio label:hover { color: #ffffff !important; }

/* Remove all default padding */
.block-container { padding: 0 2rem 2rem 2rem !important; }

/* Page header */
.page-header {
    padding: 28px 0 20px 0;
    margin-bottom: 28px;
    border-bottom: 1px solid #1e2130;
}
.page-title {
    font-size: 1.45rem;
    font-weight: 600;
    color: #ffffff;
    letter-spacing: -0.02em;
    margin: 0;
}
.page-subtitle {
    font-size: 0.82rem;
    color: #4a4f6a;
    margin-top: 4px;
}

/* Logo */
.logo {
    font-size: 1.1rem;
    font-weight: 700;
    color: #ffffff;
    letter-spacing: -0.02em;
}
.logo-accent { color: #6366f1; }
.logo-sub {
    font-size: 0.72rem;
    color: #3a3f58;
    margin-top: 2px;
    text-transform: uppercase;
    letter-spacing: 0.06em;
}

/* KPI Cards */
.kpi-grid { display: flex; gap: 14px; margin-bottom: 28px; }
.kpi-card {
    flex: 1;
    background: #111318;
    border: 1px solid #1e2130;
    border-radius: 10px;
    padding: 18px 20px;
}
.kpi-val {
    font-size: 1.8rem;
    font-weight: 700;
    color: #ffffff;
    letter-spacing: -0.03em;
    line-height: 1;
}
.kpi-lbl {
    font-size: 0.72rem;
    color: #4a4f6a;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    margin-top: 8px;
}
.kpi-sub-good { font-size: 0.75rem; color: #4ade80; margin-top: 5px; }
.kpi-sub-bad  { font-size: 0.75rem; color: #f87171; margin-top: 5px; }
.kpi-sub-neu  { font-size: 0.75rem; color: #6366f1; margin-top: 5px; }

/* Section label */
.sec-label {
    font-size: 0.75rem;
    font-weight: 600;
    color: #4a4f6a;
    text-transform: uppercase;
    letter-spacing: 0.09em;
    margin-bottom: 14px;
    padding-bottom: 8px;
    border-bottom: 1px solid #1e2130;
}

/* Chat */
.chat-wrap { display: flex; flex-direction: column; gap: 20px; padding-bottom: 16px; }

.msg-user-wrap { display: flex; justify-content: flex-end; }
.msg-user {
    background: #1a1f35;
    border: 1px solid #252b45;
    border-radius: 14px 14px 4px 14px;
    padding: 12px 16px;
    max-width: 72%;
    font-size: 0.9rem;
    color: #c9cde0;
    line-height: 1.55;
}

.msg-bot-wrap { display: flex; flex-direction: column; gap: 6px; }
.msg-bot-label {
    font-size: 0.7rem;
    font-weight: 600;
    color: #6366f1;
    text-transform: uppercase;
    letter-spacing: 0.08em;
}
.msg-bot {
    background: #111318;
    border: 1px solid #1e2130;
    border-radius: 4px 14px 14px 14px;
    padding: 16px 20px;
    max-width: 88%;
    font-size: 0.9rem;
    color: #c9cde0;
    line-height: 1.75;
}
.msg-bot p { margin: 0 0 12px 0; }
.msg-bot p:last-child { margin-bottom: 0; }

/* Suggest chips */
.chip-wrap { display: flex; flex-wrap: wrap; gap: 8px; margin-bottom: 24px; }
.chip {
    background: #111318;
    border: 1px solid #1e2130;
    border-radius: 20px;
    padding: 6px 14px;
    font-size: 0.8rem;
    color: #8b90a8;
    cursor: pointer;
}
.chip:hover { border-color: #6366f1; color: #c9cde0; }

/* Divider */
.divider { border: none; border-top: 1px solid #1e2130; margin: 24px 0; }

/* Metric overrides */
[data-testid="metric-container"] {
    background: #111318;
    border: 1px solid #1e2130;
    border-radius: 10px;
    padding: 16px 18px;
}
[data-testid="metric-container"] label { color: #4a4f6a !important; font-size: 0.72rem !important; }
[data-testid="metric-container"] [data-testid="stMetricValue"] { color: #ffffff !important; font-size: 1.6rem !important; }

/* Table */
.stDataFrame { border-radius: 10px; overflow: hidden; }
.stDataFrame thead tr th {
    background: #111318 !important;
    color: #4a4f6a !important;
    font-size: 0.72rem !important;
    text-transform: uppercase;
    letter-spacing: 0.07em;
}

/* Inputs */
.stTextInput input, .stSelectbox > div, .stMultiSelect > div {
    background: #111318 !important;
    border: 1px solid #1e2130 !important;
    border-radius: 8px !important;
    color: #c9cde0 !important;
}

/* Download button */
.stDownloadButton button {
    background: #6366f1 !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 8px !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    padding: 8px 16px !important;
}

/* Clear button */
.stButton button {
    background: #111318 !important;
    color: #6b7094 !important;
    border: 1px solid #1e2130 !important;
    border-radius: 8px !important;
    font-size: 0.82rem !important;
}

/* Hide streamlit chrome */
#MainMenu, footer, header { visibility: hidden; }
.stDeployButton { display: none; }
</style>
""", unsafe_allow_html=True)


# ── Data ──────────────────────────────────────────────────────────────────────
@st.cache_data
def load_data(file=None):
    source = file if file is not None else "branch_atm_report.xlsx"
    try:
        mo = pd.read_excel(source, sheet_name="Monthly_Data")
        bs = pd.read_excel(source, sheet_name="Branch_Summary")
        fl = pd.read_excel(source, sheet_name="Flagged_Branches")
        rs = pd.read_excel(source, sheet_name="Region_Summary")
        mo["Month"] = pd.to_datetime(mo["Month"])
        fl["Month"] = pd.to_datetime(fl["Month"])
        return mo, bs, fl, rs
    except FileNotFoundError:
        st.error("branch_atm_report.xlsx not found. Run pipeline.py first.")
        st.stop()

REGIONS_DEFAULT = ["All Regions"]

# Plotly base theme
CHART = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    font=dict(color="#4a4f6a", family="Inter", size=11),
    margin=dict(l=0, r=0, t=32, b=0),
)
AXIS = dict(gridcolor="#1a1d28", linecolor="#1e2130", tickcolor="#1e2130")

def chart(**kwargs):
    """Build a layout dict from CHART base + per-chart overrides, always injecting clean axes."""
    layout = {**CHART}
    if "xaxis" not in kwargs:
        layout["xaxis"] = AXIS
    if "yaxis" not in kwargs:
        layout["yaxis"] = AXIS
    layout.update(kwargs)
    return layout
PAL = ["#6366f1","#22d3ee","#4ade80","#f59e0b","#f87171"]


# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
        <div style='padding:16px 0 24px 0;'>
            <div class='logo'>Branch <span class='logo-accent'>Copilot</span></div>
            <div class='logo-sub'>Branch & ATM Intelligence</div>
        </div>
    """, unsafe_allow_html=True)

    page = st.radio("", ["Chat", "Analytics", "Flagged Branches"], label_visibility="collapsed")

    st.markdown("<hr style='border-color:#1e2130; margin:16px 0;'>", unsafe_allow_html=True)

    api_key = st.text_input("Groq API Key", type="password",
                             placeholder="Paste your key here",
                             help="Required for the Chat page only")

    st.markdown("<hr style='border-color:#1e2130; margin:16px 0;'>", unsafe_allow_html=True)

    uploaded_file = st.file_uploader("Upload Client Data (.xlsx)", type=["xlsx"],
                                      help="Upload a branch_atm_report.xlsx file to analyze a different client")
    if uploaded_file:
        st.session_state["uploaded_file"] = uploaded_file
        st.success("File loaded")

    st.markdown("<hr style='border-color:#1e2130; margin:16px 0;'>", unsafe_allow_html=True)

    selected_region = st.selectbox("Region", REGIONS_DEFAULT)

    st.markdown("""
        <div style='font-size:0.7rem; color:#2a2f48; margin-top:24px; line-height:1.8;'>
            75 branches &nbsp;·&nbsp; 5 regions<br>
            Jan 2023 – Jun 2024<br>
            18 months of performance data
        </div>
    """, unsafe_allow_html=True)


# Load data after sidebar so uploaded_file is in session_state
uploaded = st.session_state.get("uploaded_file", None)
monthly, branch_summary, flagged, region_summary = load_data(uploaded)
REGIONS = ["All Regions"] + sorted(branch_summary["Region"].unique().tolist())

# Update the region filter now that we have real region names
with st.sidebar:
    selected_region = st.selectbox("Region", REGIONS, key="region_select")

def filter_region(df, col="Region"):
    return df if selected_region == "All Regions" else df[df[col] == selected_region]


# ══════════════════════════════════════════════════════════════════════════════
# CHAT PAGE
# ══════════════════════════════════════════════════════════════════════════════
if page == "Chat":
    st.markdown("""
        <div class='page-header'>
            <div class='page-title'>Chat</div>
            <div class='page-subtitle'>Ask anything about your branch and ATM network in plain English</div>
        </div>
    """, unsafe_allow_html=True)

    if "messages" not in st.session_state:
        st.session_state.messages     = []
        st.session_state.last_summary = None
        st.session_state.last_q       = None

    # Suggested questions (only when chat is empty)
    if not st.session_state.messages:
        st.markdown("<div class='sec-label'>Suggested questions</div>", unsafe_allow_html=True)
        suggestions = [
            "Give me an overview of the Northeast region",
            "Which branches are struggling the most?",
            "Summarize ATM performance across all regions",
            "What is happening in the Southwest?",
            "Which regions have the best customer satisfaction?",
            "Show me branches with declining transactions",
        ]
        cols = st.columns(3)
        for i, s in enumerate(suggestions):
            if cols[i % 3].button(s, key=f"s{i}", use_container_width=True):
                st.session_state["_prefill"] = s
        st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    # Render chat history
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(f"""
                <div class='msg-user-wrap'>
                    <div class='msg-user'>{msg['content']}</div>
                </div>""", unsafe_allow_html=True)
        else:
            # Format paragraphs cleanly
            paras = [p.strip() for p in msg["content"].split("\n\n") if p.strip()]
            body  = "".join(f"<p>{p}</p>" for p in paras)
            st.markdown(f"""
                <div class='msg-bot-wrap'>
                    <div class='msg-bot-label'>Branch Copilot</div>
                    <div class='msg-bot'>{body}</div>
                </div>""", unsafe_allow_html=True)

    # Build data context
    def build_context(q):
        ql = q.lower()
        region_list = ["northeast","southeast","midwest","southwest","west"]
        target = next((r.title() for r in region_list if r in ql), None)
        if selected_region != "All Regions":
            target = selected_region

        bs = filter_region(branch_summary)
        lines = [
            "=== NETWORK SNAPSHOT ===",
            f"Total branches: {len(bs)}",
            f"Flagged: {int(bs['Is_Flagged'].sum())} ({bs['Is_Flagged'].mean()*100:.0f}%)",
            f"Avg ATM uptime: {bs['ATM_Uptime_Pct'].mean()*100:.1f}%",
            f"Avg CSAT: {bs['CSAT_Score'].mean():.1f}/100",
            f"Avg ROI: {bs['ROI_Pct'].mean():.1f}%", ""
        ]
        if target:
            rb = branch_summary[branch_summary["Region"] == target]
            lines += [
                f"=== {target.upper()} ===",
                f"Branches: {len(rb)}, Flagged: {int(rb['Is_Flagged'].sum())}",
                f"Avg transactions/month: {rb['Transaction_Volume'].mean():,.0f}",
                f"Avg ATM uptime: {rb['ATM_Uptime_Pct'].mean()*100:.1f}%",
                f"Avg CSAT: {rb['CSAT_Score'].mean():.1f}",
                f"Avg ROI: {rb['ROI_Pct'].mean():.1f}%", "",
            ]
            rf = flagged[flagged["Region"]==target][
                ["Branch_ID","Month","Transaction_Volume","ATM_Uptime_Pct","CSAT_Score","Flag_Reasons"]
            ].tail(12)
            if len(rf):
                lines += ["Flagged records:", rf.to_string(index=False)]
        elif any(w in ql for w in ["atm","uptime"]):
            atm = bs[bs["ATM_Uptime_Pct"]<0.92].sort_values("ATM_Uptime_Pct")
            lines += ["=== ATM ISSUES ===",
                      atm[["Branch_ID","Region","ATM_Uptime_Pct","Flag_Reasons"]].head(10).to_string(index=False)]
        elif any(w in ql for w in ["struggling","worst","flag","problem","decline"]):
            top = bs[bs["Is_Flagged"]].sort_values("Total_Flag_Months", ascending=False)
            lines += ["=== TOP FLAGGED ===",
                      top[["Branch_ID","Region","Branch_Type","Transaction_Volume",
                            "ATM_Uptime_Pct","CSAT_Score","ROI_Pct",
                            "Total_Flag_Months","Flag_Reasons"]].head(12).to_string(index=False)]
        else:
            latest = region_summary.groupby("Region").last().reset_index()
            lines += ["=== REGIONAL BREAKDOWN ===",
                      latest[["Region","Total_Transactions","Avg_ATM_Uptime_Pct",
                               "Avg_CSAT_Score","Branches_Flagged","Total_Branches"]].to_string(index=False)]
        return "\n".join(lines)

    def call_groq(question, context):
        client = Groq(api_key=api_key)
        system = (
            "You are a senior analyst at a bank advisory firm answering questions from a bank executive. "
            "Be extremely concise. Maximum 4 short sentences total. "
            "Lead with the single most important number or finding. "
            "Then one sentence on what is working. "
            "Then one sentence on what needs attention. "
            "End with one clear action to take. "
            "Plain English only. No jargon. No bullet points. No headers. No long paragraphs."
        )
        r = Groq(api_key=api_key).chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role":"system","content":system},
                {"role":"user","content":f'Question: "{question}"\n\nData:\n{context}'},
            ],
            temperature=0.4, max_tokens=800,
        )
        return r.choices[0].message.content.strip()

    def make_docx(question, summary):
        doc = Document()
        for s in doc.sections:
            s.top_margin = s.bottom_margin = Inches(1.0)
            s.left_margin = s.right_margin = Inches(1.2)
        h = doc.add_paragraph()
        r = h.add_run("BRANCH COPILOT  |  Advisory Intelligence")
        r.font.size = Pt(9); r.bold = True
        r.font.color.rgb = RGBColor(0x63,0x66,0xf1)
        doc.add_paragraph()
        t = doc.add_heading("Branch & ATM Performance Summary", 1)
        for run in t.runs:
            run.font.color.rgb = RGBColor(0x1F,0x38,0x64)
            run.font.size = Pt(18)
        m = doc.add_paragraph()
        m.add_run(f"Generated: {datetime.today().strftime('%B %d, %Y')}  |  Data through: June 2024").font.size = Pt(9)
        doc.add_paragraph("─"*72)
        doc.add_paragraph()
        qp = doc.add_paragraph(); qp.add_run("Question: ").bold = True; qp.add_run(question)
        doc.add_paragraph()
        for para in summary.split("\n\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                for run in p.runs: run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(10)
        doc.add_paragraph()
        foot = doc.add_paragraph()
        fr = foot.add_run("AI-generated from branch performance data. Validate before client delivery.")
        fr.font.size = Pt(8); fr.italic = True
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    # Input
    prefill = st.session_state.pop("_prefill", None)
    question = st.chat_input("Ask about your branch network...") or prefill

    if question:
        if not api_key:
            st.warning("Enter your Groq API key in the sidebar to use the chat.")
        else:
            st.session_state.messages.append({"role":"user","content":question})
            with st.spinner(""):
                ctx     = build_context(question)
                summary = call_groq(question, ctx)
            st.session_state.messages.append({"role":"assistant","content":summary})
            st.session_state.last_summary = summary
            st.session_state.last_q       = question
            st.rerun()

    if st.session_state.get("last_summary"):
        st.markdown("<hr class='divider'>", unsafe_allow_html=True)
        c1, c2 = st.columns([5,1])
        with c2:
            buf = make_docx(st.session_state.last_q, st.session_state.last_summary)
            st.download_button(
                "Download summary",
                data=buf,
                file_name=f"branch_summary_{datetime.today().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with c1:
            if st.button("Clear conversation"):
                st.session_state.messages     = []
                st.session_state.last_summary = None
                st.session_state.last_q       = None
                st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# ANALYTICS PAGE
# ══════════════════════════════════════════════════════════════════════════════
elif page == "Analytics":
    st.markdown("""
        <div class='page-header'>
            <div class='page-title'>Analytics</div>
            <div class='page-subtitle'>Network-wide performance dashboard · Jan 2023 – Jun 2024</div>
        </div>
    """, unsafe_allow_html=True)

    bs = filter_region(branch_summary)
    mo = filter_region(monthly)
    rs = filter_region(region_summary)

    # ── KPI row ───────────────────────────────────────────────────────────────
    k = st.columns(5)
    def kpi(col, val, lbl, sub=None, sub_type="neu"):
        sub_html = f"<div class='kpi-sub-{sub_type}'>{sub}</div>" if sub else ""
        col.markdown(f"""
            <div class='kpi-card'>
                <div class='kpi-val'>{val}</div>
                <div class='kpi-lbl'>{lbl}</div>
                {sub_html}
            </div>""", unsafe_allow_html=True)

    kpi(k[0], len(bs),                                  "Branches Monitored")
    kpi(k[1], int(bs["Is_Flagged"].sum()),              "Branches Flagged",
        f"{bs['Is_Flagged'].mean()*100:.0f}% of network", "bad")
    kpi(k[2], f"{bs['ATM_Uptime_Pct'].mean()*100:.1f}%","Avg ATM Uptime",
        "Above 90% threshold", "good")
    kpi(k[3], f"{bs['CSAT_Score'].mean():.1f}",         "Avg CSAT Score", "Out of 100")
    kpi(k[4], f"{bs['ROI_Pct'].mean():.1f}%",           "Avg Network ROI")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Row 1: Transaction trend (full width) ─────────────────────────────────
    st.markdown("<div class='sec-label'>Transaction Volume Trend by Region</div>", unsafe_allow_html=True)
    trend = mo.groupby(["Month","Region"])["Transaction_Volume"].sum().reset_index()
    fig1  = px.line(trend, x="Month", y="Transaction_Volume", color="Region",
                    color_discrete_sequence=PAL,
                    labels={"Transaction_Volume":"Total Transactions","Month":""})
    fig1.update_traces(line=dict(width=2.5))
    fig1.update_layout(**chart(height=260,
        legend=dict(orientation="h", y=1.08, x=0, font=dict(color="#6b7094", size=11), bgcolor="rgba(0,0,0,0)")))
    st.plotly_chart(fig1, use_container_width=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Row 2: Branch health donut + Flagged by region ────────────────────────
    c1, spacer, c2 = st.columns([4, 0.4, 6])

    with c1:
        st.markdown("<div class='sec-label'>Branch Health Overview</div>", unsafe_allow_html=True)
        healthy  = len(bs) - int(bs["Is_Flagged"].sum())
        flagged_n = int(bs["Is_Flagged"].sum())
        fig2 = go.Figure(go.Pie(
            labels=["Healthy", "Needs Attention"],
            values=[healthy, flagged_n],
            hole=0.65,
            marker=dict(colors=["#4ade80","#f87171"],
                        line=dict(color="#0d0f14", width=3)),
            textinfo="none",
            hovertemplate="%{label}: %{value} branches<extra></extra>",
        ))
        fig2.add_annotation(text=f"<b>{flagged_n}</b><br><span style='font-size:11px'>flagged</span>",
                            x=0.5, y=0.5, showarrow=False,
                            font=dict(size=20, color="#ffffff", family="Inter"))
        fig2.update_layout(**chart(height=240, showlegend=True,
                           legend=dict(orientation="h", y=-0.1, x=0.15,
                                       font=dict(color="#6b7094", size=11),
                                       bgcolor="rgba(0,0,0,0)")))
        st.plotly_chart(fig2, use_container_width=True)

    with c2:
        st.markdown("<div class='sec-label'>Flagged Branches by Region</div>", unsafe_allow_html=True)
        reg_flags = bs.groupby("Region").agg(
            Flagged=("Is_Flagged","sum"),
            Total=("Branch_ID","count")
        ).reset_index()
        reg_flags["Healthy"] = reg_flags["Total"] - reg_flags["Flagged"]
        reg_flags = reg_flags.sort_values("Flagged", ascending=True)
        fig3 = go.Figure()
        fig3.add_trace(go.Bar(name="Healthy", y=reg_flags["Region"],
                              x=reg_flags["Healthy"], orientation="h",
                              marker_color="#4ade80",
                              hovertemplate="%{y}: %{x} healthy branches<extra></extra>"))
        fig3.add_trace(go.Bar(name="Flagged", y=reg_flags["Region"],
                              x=reg_flags["Flagged"], orientation="h",
                              marker_color="#f87171", opacity=0.85,
                              hovertemplate="%{y}: %{x} flagged branches<extra></extra>"))
        fig3.update_layout(**chart(barmode="stack", height=240,
                           legend=dict(orientation="h", y=1.12, x=0,
                                       font=dict(color="#6b7094", size=11),
                                       bgcolor="rgba(0,0,0,0)"),
                           xaxis=dict(gridcolor="#1a1d28", linecolor="#1e2130",
                                      title="Number of Branches")))
        st.plotly_chart(fig3, use_container_width=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Row 3: Top struggling branches + CSAT trend ───────────────────────────
    c3, spacer2, c4 = st.columns([5, 0.4, 5])

    with c3:
        st.markdown("<div class='sec-label'>Top 10 Branches Needing Attention</div>",
                    unsafe_allow_html=True)
        top10 = bs[bs["Is_Flagged"]].sort_values("Total_Flag_Months", ascending=False).head(10)
        top10["label"] = top10["Branch_ID"] + " (" + top10["Region"].str[:2] + ")"
        fig4 = px.bar(top10, x="Total_Flag_Months", y="label", orientation="h",
                      color="Total_Flag_Months",
                      color_continuous_scale=[[0,"#fbbf24"],[1,"#f87171"]],
                      labels={"Total_Flag_Months":"Months Flagged","label":""},
                      hover_data={"Flag_Reasons":True,"CSAT_Score":True,
                                  "ATM_Uptime_Pct":True,"label":False,
                                  "Total_Flag_Months":True})
        fig4.update_layout(**chart(height=320, coloraxis_showscale=False,
                           yaxis=dict(gridcolor="rgba(0,0,0,0)", linecolor="#1e2130", autorange="reversed")))
        st.plotly_chart(fig4, use_container_width=True)

    with c4:
        st.markdown("<div class='sec-label'>Network CSAT Score Over Time</div>",
                    unsafe_allow_html=True)
        csat_trend = mo.groupby("Month")["CSAT_Score"].mean().reset_index()
        fig5 = px.line(csat_trend, x="Month", y="CSAT_Score",
                       labels={"CSAT_Score":"Avg CSAT Score","Month":""},
                       color_discrete_sequence=["#6366f1"])
        fig5.update_traces(line=dict(width=2.5))
        fig5.add_hline(y=65, line_dash="dot", line_color="#f87171",
                       annotation_text="Alert threshold (65)",
                       annotation_font=dict(color="#f87171", size=10))
        fig5.add_hrect(y0=0, y1=65, fillcolor="rgba(248,113,113,0.05)",
                       line_width=0)
        fig5.update_layout(**chart(height=320,
            yaxis=dict(gridcolor="#1a1d28", linecolor="#1e2130", range=[60, 85])))
        st.plotly_chart(fig5, use_container_width=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Row 4: ATM problem branches ───────────────────────────────────────────
    st.markdown("<div class='sec-label'>ATM Uptime — Problem Branches (Below 90%)</div>",
                unsafe_allow_html=True)
    atm_problems = bs[bs["ATM_Uptime_Pct"] < 0.92].copy()
    atm_problems["Uptime_Pct"] = (atm_problems["ATM_Uptime_Pct"]*100).round(1)
    atm_problems["label"] = atm_problems["Branch_ID"] + "  ·  " + atm_problems["Region"] + "  ·  " + atm_problems["Branch_Type"]
    atm_problems = atm_problems.sort_values("Uptime_Pct")

    if len(atm_problems):
        fig6 = px.bar(atm_problems, x="Uptime_Pct", y="label", orientation="h",
                      color="Uptime_Pct",
                      color_continuous_scale=[[0,"#f87171"],[0.5,"#fbbf24"],[1,"#facc15"]],
                      range_color=[75, 92],
                      labels={"Uptime_Pct":"ATM Uptime %","label":""},
                      hover_data={"CSAT_Score":True,"Flag_Reasons":True,
                                  "Uptime_Pct":True,"label":False})
        fig6.add_vline(x=90, line_dash="dot", line_color="#f87171",
                       annotation_text="90% minimum",
                       annotation_font=dict(color="#f87171", size=10))
        fig6.update_layout(**chart(height=max(200, len(atm_problems)*48),
                           coloraxis_showscale=False,
                           yaxis=dict(gridcolor="rgba(0,0,0,0)", linecolor="#1e2130"),
                           xaxis=dict(gridcolor="#1a1d28", linecolor="#1e2130",
                                      range=[75, 100], title="ATM Uptime %")))
        st.plotly_chart(fig6, use_container_width=True)
    else:
        st.success("No ATM uptime issues detected in the selected region.")


# ══════════════════════════════════════════════════════════════════════════════
# FLAGGED BRANCHES PAGE
# ══════════════════════════════════════════════════════════════════════════════
elif page == "Flagged Branches":
    st.markdown("""
        <div class='page-header'>
            <div class='page-title'>Flagged Branches</div>
            <div class='page-subtitle'>Branches automatically identified as needing attention</div>
        </div>
    """, unsafe_allow_html=True)

    bs_f = filter_region(branch_summary)
    fo   = bs_f[bs_f["Is_Flagged"]].copy()

    # Metrics
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Flagged Branches",     int(fo["Is_Flagged"].sum()))
    m2.metric("ATM Issues",           int((fo["ATM_Uptime_Pct"] < 0.90).sum()))
    m3.metric("Low CSAT",             int((fo["CSAT_Score"] < 65).sum()))
    m4.metric("Avg Uptime (flagged)", f"{fo['ATM_Uptime_Pct'].mean()*100:.1f}%")

    st.markdown("<hr class='divider'>", unsafe_allow_html=True)

    # Filters
    f1, f2 = st.columns([3,1])
    with f1:
        issue_filter = st.multiselect(
            "Filter by issue",
            ["Transaction decline >15%","ATM uptime <90%","CSAT below 65",
             "CSAT dropped 10+ pts","Negative ROI"],
            placeholder="Show all issues"
        )
    with f2:
        sort_by = st.selectbox("Sort by", ["Total_Flag_Months","ATM_Uptime_Pct",
                                            "CSAT_Score","ROI_Pct","Transaction_Volume"])

    if issue_filter:
        fo = fo[fo["Flag_Reasons"].apply(lambda r: any(i in str(r) for i in issue_filter))]

    fo = fo.sort_values(sort_by, ascending=(sort_by == "ATM_Uptime_Pct"))

    st.markdown(f"<div class='sec-label'>Showing {len(fo)} branches</div>", unsafe_allow_html=True)

    display_cols = ["Branch_ID","Region","State","Branch_Type",
                    "Transaction_Volume","ATM_Uptime_Pct","CSAT_Score",
                    "ROI_Pct","Total_Flag_Months","Flag_Reasons"]
    disp = fo[display_cols].copy()
    disp["ATM_Uptime_Pct"] = (disp["ATM_Uptime_Pct"]*100).round(1).astype(str)+"%"
    disp["ROI_Pct"]        = disp["ROI_Pct"].round(1).astype(str)+"%"
    disp = disp.rename(columns={
        "Branch_ID":"Branch","Transaction_Volume":"Transactions",
        "ATM_Uptime_Pct":"ATM Uptime","CSAT_Score":"CSAT",
        "ROI_Pct":"ROI","Total_Flag_Months":"Months Flagged","Flag_Reasons":"Issues"
    })

    st.dataframe(disp, use_container_width=True, height=480, hide_index=True)

    st.markdown("<br>", unsafe_allow_html=True)
    csv = fo[display_cols].to_csv(index=False)
    st.download_button(
        "Export as CSV",
        data=csv,
        file_name=f"flagged_branches_{datetime.today().strftime('%Y%m%d')}.csv",
        mime="text/csv",
    )
