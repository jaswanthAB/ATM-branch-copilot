"""
Phase 4: AI Market Summary Tool
Diebold Nixdorf Advisory Services — Case Study Project

Usage:
    python ai_summary.py

You will be prompted to type a question like:
    "Give me an overview of the Northeast region"
    "Which branches are struggling the most?"
    "Summarize ATM performance across all regions"

The tool pulls relevant data from branch_atm_report.xlsx,
sends it to Groq's LLM, and saves a client-ready Word document.

Requirements:
    pip install groq openpyxl python-docx pandas
    Set your Groq API key when prompted (or set GROQ_API_KEY env variable)
"""

import os
import sys
import pandas as pd
from datetime import datetime
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Load data ─────────────────────────────────────────────────────────────────
def load_data(path="branch_atm_report.xlsx"):
    try:
        monthly        = pd.read_excel(path, sheet_name="Monthly_Data")
        branch_summary = pd.read_excel(path, sheet_name="Branch_Summary")
        flagged        = pd.read_excel(path, sheet_name="Flagged_Branches")
        region_summary = pd.read_excel(path, sheet_name="Region_Summary")
        return monthly, branch_summary, flagged, region_summary
    except FileNotFoundError:
        print("ERROR: branch_atm_report.xlsx not found.")
        print("Run pipeline.py first to generate the report file.")
        sys.exit(1)


# ── Build context from data based on the question ────────────────────────────
def build_data_context(question: str, monthly, branch_summary, flagged, region_summary) -> str:
    q = question.lower()

    # Detect region filter
    regions = ["northeast", "southeast", "midwest", "southwest", "west"]
    target_region = next((r.title() for r in regions if r in q), None)

    lines = []

    # ── Overall snapshot ──────────────────────────────────────────────────────
    lines.append("=== NETWORK OVERVIEW (Latest Month) ===")
    latest_month = branch_summary.copy()
    lines.append(f"Total branches: {len(latest_month)}")
    lines.append(f"Branches flagged for issues: {latest_month['Is_Flagged'].sum()} ({latest_month['Is_Flagged'].mean()*100:.0f}%)")
    lines.append(f"Average ATM uptime: {latest_month['ATM_Uptime_Pct'].mean()*100:.1f}%")
    lines.append(f"Average CSAT score: {latest_month['CSAT_Score'].mean():.1f}/100")
    lines.append(f"Average ROI: {latest_month['ROI_Pct'].mean():.1f}%")
    lines.append("")

    # ── Region filter if applicable ───────────────────────────────────────────
    if target_region:
        lines.append(f"=== {target_region.upper()} REGION DETAIL ===")
        reg_branches = branch_summary[branch_summary["Region"] == target_region]
        lines.append(f"Branches in region: {len(reg_branches)}")
        lines.append(f"Flagged branches: {reg_branches['Is_Flagged'].sum()}")
        lines.append(f"Avg transaction volume: {reg_branches['Transaction_Volume'].mean():,.0f}/month")
        lines.append(f"Avg ATM uptime: {reg_branches['ATM_Uptime_Pct'].mean()*100:.1f}%")
        lines.append(f"Avg CSAT: {reg_branches['CSAT_Score'].mean():.1f}")
        lines.append(f"Avg ROI: {reg_branches['ROI_Pct'].mean():.1f}%")
        lines.append("")

        # Flagged branches in this region
        reg_flagged = flagged[flagged["Region"] == target_region][
            ["Branch_ID", "Month", "Transaction_Volume", "ATM_Uptime_Pct", "CSAT_Score", "Flag_Reasons"]
        ].tail(20)
        if len(reg_flagged) > 0:
            lines.append(f"Flagged branch records ({target_region}):")
            lines.append(reg_flagged.to_string(index=False))
        lines.append("")

        # Region trend (last 6 months)
        reg_trend = region_summary[region_summary["Region"] == target_region].tail(6)
        lines.append(f"Monthly trend ({target_region}, last 6 months):")
        lines.append(reg_trend[["Period","Total_Transactions","Avg_ATM_Uptime_Pct","Avg_CSAT_Score","Branches_Flagged"]].to_string(index=False))

    else:
        # ── All-regions summary ───────────────────────────────────────────────
        lines.append("=== REGIONAL BREAKDOWN (Latest Period) ===")
        latest_period = region_summary.groupby("Region").last().reset_index()
        lines.append(latest_period[[
            "Region","Total_Transactions","Avg_ATM_Uptime_Pct",
            "Avg_CSAT_Score","Branches_Flagged","Total_Branches","ROI_Pct"
        ]].to_string(index=False))
        lines.append("")

        # ── Top flagged branches ──────────────────────────────────────────────
        if "atm" in q or "uptime" in q:
            atm_issues = branch_summary[branch_summary["ATM_Uptime_Pct"] < 0.92].sort_values("ATM_Uptime_Pct")
            lines.append("=== BRANCHES WITH ATM ISSUES ===")
            lines.append(atm_issues[["Branch_ID","Region","ATM_Uptime_Pct","CSAT_Score","Flag_Reasons"]].head(10).to_string(index=False))

        elif "struggling" in q or "worst" in q or "flag" in q or "problem" in q:
            top_flagged = branch_summary[branch_summary["Is_Flagged"]].sort_values("Total_Flag_Months", ascending=False)
            lines.append("=== MOST FLAGGED BRANCHES ===")
            lines.append(top_flagged[["Branch_ID","Region","Branch_Type","Transaction_Volume",
                                       "ATM_Uptime_Pct","CSAT_Score","ROI_Pct",
                                       "Total_Flag_Months","Flag_Reasons"]].head(12).to_string(index=False))

        else:
            # General: show branch summary top 15
            lines.append("=== BRANCH PERFORMANCE SNAPSHOT (Top Flagged) ===")
            top = branch_summary[branch_summary["Is_Flagged"]].head(15)
            lines.append(top[["Branch_ID","Region","Branch_Type","Transaction_Volume",
                               "ATM_Uptime_Pct","CSAT_Score","ROI_Pct","Flag_Reasons"]].to_string(index=False))

    return "\n".join(lines)


# ── Call Groq API ─────────────────────────────────────────────────────────────
def call_groq(api_key: str, question: str, data_context: str) -> str:
    client = Groq(api_key=api_key)

    system_prompt = """You are a senior analyst at Diebold Nixdorf Advisory Services.
Your job is to turn branch and ATM performance data into clear, concise written summaries
for bank executives and regional managers who are NOT technical.

Rules:
- Write in plain English. No jargon, no technical terms.
- Use short paragraphs (3-5 sentences each).
- Lead with the most important finding.
- Always include: what's going well, what needs attention, and one clear recommendation.
- Reference specific numbers from the data (branch counts, percentages, scores).
- Tone: professional but direct. Like a trusted advisor, not a data report.
- Length: 3-4 paragraphs. No bullet points. No headers inside the summary.
- End with one sentence that frames the next step for the client."""

    user_prompt = f"""The client asked: "{question}"

Here is the relevant data:

{data_context}

Write a client-ready summary based on this data. Remember: the reader is a bank executive or regional manager, not a data analyst."""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        temperature=0.4,
        max_tokens=800,
    )

    return response.choices[0].message.content.strip()


# ── Save as Word document ─────────────────────────────────────────────────────
def save_word_doc(question: str, summary_text: str, region: str = None) -> str:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Header bar ────────────────────────────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run("DIEBOLD NIXDORF  |  Advisory Services")
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.bold  = True

    doc.add_paragraph()  # spacer

    # ── Title ─────────────────────────────────────────────────────────────────
    region_label = f"{region} Region — " if region else ""
    title_text   = f"{region_label}Branch & ATM Performance Overview"
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(18)

    # ── Date and prepared by ──────────────────────────────────────────────────
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(
        f"Prepared: {datetime.today().strftime('%B %d, %Y')}    |    "
        f"Data through: June 2024    |    Prepared by: Advisory Services Analytics"
    )
    meta_run.font.size  = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Divider ───────────────────────────────────────────────────────────────
    doc.add_paragraph("─" * 72)

    doc.add_paragraph()  # spacer

    # ── Client question ───────────────────────────────────────────────────────
    q_label = doc.add_paragraph()
    q_run = q_label.add_run("Client Request: ")
    q_run.bold = True
    q_run.font.size = Pt(10)
    q_label.add_run(question).font.size = Pt(10)

    doc.add_paragraph()

    # ── Summary body ──────────────────────────────────────────────────────────
    for para_text in summary_text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph(para_text)
        p.style.font.size = Pt(11)
        for run in p.runs:
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "This summary was generated from the branch and ATM performance dataset. "
        "Data is simulated for demonstration purposes. "
        "Findings should be validated against live client data before client delivery."
    )
    footer_run.font.size  = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    footer_run.font.italic = True

    # ── Save ──────────────────────────────────────────────────────────────────
    region_slug  = region.lower().replace(" ", "_") + "_" if region else "network_"
    timestamp    = datetime.today().strftime("%Y%m%d_%H%M")
    filename     = f"summary_{region_slug}{timestamp}.docx"
    doc.save(filename)
    return filename


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("  Diebold Nixdorf — AI Branch Performance Summary Tool")
    print("=" * 60)
    print()

    # API key
    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if not api_key:
        api_key = input("Enter your Groq API key: ").strip()
    if not api_key:
        print("No API key provided. Exiting.")
        sys.exit(1)

    # Load data
    monthly, branch_summary, flagged, region_summary = load_data()
    print(f"Data loaded: {len(branch_summary)} branches, {len(flagged)} flagged records\n")

    # Question
    print("Example questions:")
    print("  - Give me an overview of the Northeast region")
    print("  - Which branches are struggling the most?")
    print("  - Summarize ATM performance across all regions")
    print("  - What's happening in the Southwest?")
    print()
    question = input("Your question: ").strip()
    if not question:
        question = "Give me an overview of branch and ATM performance across all regions"

    print("\nPulling relevant data...")
    regions = ["northeast", "southeast", "midwest", "southwest", "west"]
    target_region = next((r.title() for r in regions if r in question.lower()), None)
    data_context = build_data_context(question, monthly, branch_summary, flagged, region_summary)

    print("Generating summary...")
    summary = call_groq(api_key, question, data_context)

    print("\n── GENERATED SUMMARY ────────────────────────────────────────")
    print(summary)
    print("─" * 60)

    filename = save_word_doc(question, summary, target_region)
    print(f"\nSaved → {filename}")
    print("\nDone. Open the .docx file to see the client-ready brief.")
